import re
from docx import Document
from docx.shared import RGBColor
from enum import Enum
import webcolors



class JiraWikiBlockTypes(Enum):
	LIST = "list"
	HEADING = "heading"
	TABLE = "table"
	CODE = "code"
	QUOTE = "quote"
	PARAGRAPH = "paragraph"


class JiraWiki2Docx():

	def __init__(self, jira_text=None, document=None):
		self.jira_text = jira_text
		self.document = document if document else Document()
		self.table_style = 'Table Grid'


	def detect_jira_block_type(jira_block):
		jira_block_trimmed = jira_block.strip()
		marker_heading = r"^h([1-6])\..*$"
		marker_list = "^[*#-]+\s.*$"
		marker_table = "^[||]+.*$"
		if re.match(marker_heading, jira_block_trimmed, re.DOTALL):
			return JiraWikiBlockTypes.HEADING
		if re.match(marker_list, jira_block_trimmed, re.DOTALL):
			return JiraWikiBlockTypes.LIST
		if re.match(marker_table, jira_block_trimmed, re.DOTALL):
			return JiraWikiBlockTypes.TABLE
		return JiraWikiBlockTypes.PARAGRAPH


	def get_jira_blocks(self):
		self.jira_blocks = re.split("\n\n", self.jira_text)


	def write_heading_jira_to_doc(self, text, doc_part_to_write_to):
		pattern_heading = r"^h([1-6])\.(.*)"
		heading_match = re.match(pattern_heading, text)
		heading_number, heading_text = heading_match.groups()
		para = doc_part_to_write_to.add_paragraph()
		para.style = f"Heading {heading_number}"
		JiraWiki2Docx.apply_jira_text_effect_to_text_in_docx(heading_text, para)


	def write_list_jira_to_doc(self, text, doc_part_to_write_to):
		pattern_list_item = r"^([\*|\-|\#]+)\s{1}(.*)"
		# split block into lines
		lines = re.split("\n", text)
		for line in lines:
			para = doc_part_to_write_to.add_paragraph()
			list_type, list_text = re.match(pattern_list_item, line).groups()
			# issue 1: couldn't resolve the round and square bullets differently
			# issue 2: find a way of decorating nested lists with different point styles
			# issue 3: sublevel list have unwanted top space to their parent list
			max_list_level = 3
			style_name = (
					"List Number" if list_type[:max_list_level][-1] == "#" else "List Bullet"
			)
			style_name = f'{style_name} {str(min(len(list_type), max_list_level)) if len(list_type) > 1 else ""}'.strip()
			para.style = style_name
			JiraWiki2Docx.apply_jira_text_effect_to_text_in_docx(list_text, para)


	@staticmethod
	def delete_paragraph(paragraph):
		p = paragraph._element
		p.getparent().remove(p)
		p._p = p._element = None


	def write_table_jira_to_doc(self, text, doc_part_to_write_to):
		pattern_table_row = re.compile(
			"^(\|{1,})(.*?)\|{1,}$.*?(?=^\|.*?\|$|\Z)",
			re.DOTALL | re.IGNORECASE | re.MULTILINE,
		)
		table = None
		n_cols = 0
		for row_index, row_match in enumerate(pattern_table_row.finditer(text)):
			if row_match:
				leading_markup, jira_row_inner_text = row_match.groups()
				jira_cells_text = re.split("\|{1,}", jira_row_inner_text)
				if row_index == 0:
					n_cols = len(jira_cells_text)
					table = doc_part_to_write_to.add_table(rows=1, cols=n_cols)
					table.style = self.table_style
				else:
					table.add_row()

				# write to the row columns
				for c in range(n_cols):
					JiraWiki2Docx.delete_paragraph(
						table.rows[row_index].cells[c].paragraphs[0]
					)  # remove the default paragraph in a table cell
					self.write_jira_block_to_doc(
						jira_cells_text[c], table.rows[row_index].cells[c]
					)


	def write_generic_jira_to_doc(text, doc_part_to_write_to):
		para = doc_part_to_write_to.add_paragraph()
		JiraWiki2Docx.apply_jira_text_effect_to_text_in_docx(text, para)


	def write_jira_block_to_doc(self, jira_block, doc):
		# handle leading new lines
		# actual_jira_block is jira block without leading new lines characters
		new_lines, actual_jira_block = re.match("(\n*)(.*)", jira_block, re.DOTALL).groups()
		num_new_lines = new_lines.count("\n")

		# add new lines by occurence
		for i in range(num_new_lines):
			para = doc.add_paragraph()
			para.add_run()

		actual_jira_block_type = JiraWiki2Docx.detect_jira_block_type(actual_jira_block)
		# 1. heading
		if actual_jira_block_type == JiraWikiBlockTypes.HEADING:
			self.write_heading_jira_to_doc(actual_jira_block, doc)
		# 2. list
		elif actual_jira_block_type == JiraWikiBlockTypes.LIST:
			self.write_list_jira_to_doc(actual_jira_block, doc)
		# 3. table
		elif actual_jira_block_type == JiraWikiBlockTypes.TABLE:
			# add a extea line if the previous element is table
			if self.prev_jira_block_type and self.prev_jira_block_type == JiraWikiBlockTypes.TABLE:
				doc.add_paragraph()
			self.write_table_jira_to_doc(actual_jira_block, doc)
		# 4. any other
		else:
			JiraWiki2Docx.write_generic_jira_to_doc(actual_jira_block, doc)

		self.prev_jira_block_type = actual_jira_block_type


	@staticmethod
	def detect_jira_text_effects(text):
		text_effect_para_dict = {}  # Dictionary to store detected text effects

		def detect_effects_recursive(text, parent_key=None):
			# Regular expressions for detecting various Jira text effects

			# Strong effect: Text enclosed in asterisks (*) with no space after the start or before the end.
			pattern_strong_effect = "((\*)(\S[^*]*\S)(\*))"

			# Italics effect: Text enclosed in underscores (_) with no space after the start or before the end.
			pattern_italics_effect = "((\_)(\S[^_]*\S)(\_))"

			# Deleted effect: Text enclosed in hyphens (-) with no space after the start or before the end.
			pattern_deleted_effect = "((\-)(\S[^-]*\S)(\-))"

			# Inserted effect: Text enclosed in plus signs (+) with no space after the start or before the end.
			pattern_inserted_effect = "((\+)(\S[^+]*\S)(\+))"

			# Superscript effect: Text enclosed in carets (^) with no space after the start or before the end.
			pattern_superscript_effect = "((\^)(\S[^\^]*\S)(\^))"

			# Subscript effect: Text enclosed in tildes (~) with no space after the start or before the end.
			pattern_subscript_effect = "((\~)(\S[^\~]*\S)(\~))"

			# Color effect: Text enclosed in color tags {color} with hexadecimal color code.
			pattern_color_effect = "((\{color:[#A-Za-z0-9]+\})(.*?)(\{color\}))"

			pattern_text_effect_list = [
				pattern_strong_effect,
				pattern_italics_effect,
				pattern_deleted_effect,
				pattern_inserted_effect,
				pattern_superscript_effect,
				pattern_subscript_effect,
				pattern_color_effect,
			]
			pattern_text_effect = re.compile(
				"|".join(pattern_text_effect_list), re.DOTALL | re.IGNORECASE
			)
			for m in pattern_text_effect.finditer(text):
				try:
					full_match, text_effect_ope_tag, inner_text, text_effect_cls_tag = filter(lambda x: x, m.groups())
					res = {}
					start_pos_outside_markup = m.start()
					end_pos_outside_markup = (
						m.end() - 1
					)  # decrement by 1 to match the end position in the string
					start_pos_inner_text = start_pos_outside_markup + len(
						text_effect_ope_tag
					)
					ope_tags = [text_effect_ope_tag]
					if parent_key is not None:
						res = text_effect_para_dict.pop(parent_key)
						full_match = res["full_match"]
						start_pos_outside_markup = res["start_pos_outside_markup"]
						end_pos_outside_markup = res["end_pos_outside_markup"]
						start_pos_inner_text = (
								res["start_pos_inner_text"] + start_pos_inner_text
						)
						ope_tags = res["ope_tags"] + ope_tags

					res["start_pos_outside_markup"] = start_pos_outside_markup
					res["end_pos_outside_markup"] = end_pos_outside_markup
					res["start_pos_inner_text"] = start_pos_inner_text
					res["end_pos_inner_text"] = (
							res["start_pos_inner_text"] + len(inner_text) - 1
					)
					res["inner_text"] = inner_text
					res["ope_tags"] = ope_tags
					res["full_match"] = full_match

					text_effect_para_dict[start_pos_outside_markup] = res
					detect_effects_recursive(
						inner_text, parent_key=start_pos_outside_markup
					)
				except Exception as e:
					print(e)

		detect_effects_recursive(text)
		return text_effect_para_dict

	@staticmethod
	def apply_jira_text_effect_to_text_in_docx(text, docx_element):
		detected_text_effect_list = list(JiraWiki2Docx.detect_jira_text_effects(text).values())
		prev_end_pos = -1
		if detected_text_effect_list is not None and len(detected_text_effect_list) > 0:
			for indx, item in enumerate(detected_text_effect_list):
				# apply
				inner_text = item["inner_text"]
				start_pos_outside_markup = item["start_pos_outside_markup"]
				end_pos_outside_markup = item["end_pos_outside_markup"]
				start_pos_inner_text = item["start_pos_inner_text"]
				end_pos_inner_text = item["end_pos_inner_text"]
				ope_tags = item["ope_tags"]

				leading_run_text = text[prev_end_pos + 1 : start_pos_outside_markup]
				current_run_text = text[start_pos_inner_text : end_pos_inner_text + 1]
				trailing_run_text = text[end_pos_outside_markup + 1 :]

				leading_run = docx_element.add_run(leading_run_text)
				current_run = docx_element.add_run(current_run_text)

				# apply text effects in opening tags
				for tag in ope_tags:
					if tag == "*":
						current_run.bold = True
					elif tag == "_":
						current_run.italic = True
					elif tag == "-":
						current_run.font.strike = True
					elif tag == "+":
						current_run.underline = True
					elif tag == "^":
						current_run.font.superscript = True
					elif tag == "~":
						current_run.font.subscript = True
					else:
						color_match = re.match(
								"\{color:([#A-Za-z0-9]+)\}", tag, re.IGNORECASE
						)
						if color_match:
							hex_code = color_match.groups()[0]
							if not hex_code.startswith("#"):
								hex_code = JiraWiki2Docx.color_name_to_hex(hex_code)
							if hex_code:
								current_run.font.color.rgb = RGBColor.from_string(hex_code.strip("#"))

					# handle what happens if we have gotten to the last chunk of the text effect list
					if indx == len(detected_text_effect_list) - 1:
						docx_element.add_run(trailing_run_text)

					prev_end_pos = end_pos_outside_markup
		else:
			docx_element.add_run(text)


	@staticmethod
	def color_name_to_hex(color_name):
		try:
			# Get the RGB values for the given color name
			rgb = webcolors.name_to_rgb(color_name)
			# Convert the RGB values to a hex code
			hex_code = "#{:02x}{:02x}{:02x}".format(rgb.red, rgb.green, rgb.blue)
			return hex_code
		except ValueError:
			# Handle the case when an invalid color name is provided
			print("Invalid color name.")
			return None


	def parseJira2Docx(self, save_to_file = False, output_filename = "output.docx"):
		
		# writing to docx
		self.get_jira_blocks()
		for block in self.jira_blocks:
			self.write_jira_block_to_doc(block, self.document)

		# if save_to_file is set to true, save updated document as file in the set path
		if save_to_file == True:
			self.document.save(output_filename)
		else:
			return self.document


